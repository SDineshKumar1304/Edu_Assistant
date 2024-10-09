from flask import Flask, render_template, request, send_file, session, jsonify
import google.generativeai as genai
from fpdf import FPDF
from docx import Document
import os

app = Flask(__name__)
app.secret_key = 'DK1329'  

with open('api.txt') as f:
    api_key = f.read().strip()

genai.configure(api_key=api_key)

def get_gemini_answer(question):
    """Get answer from Gemini model."""
    try:
        # Call to the Generative AI model for answering the user question
        model = genai.GenerativeModel(model_name="gemini-1.5-pro-latest")
        response = model.generate_content(question)
        
        # Ensure a clean and structured response is provided back to the user
        if response:
            return response.text.strip()  # Use the stripped, clean response
        else:
            return "Sorry, I don't have an answer for that right now."
    
    except Exception as e:
        return f"Error: {str(e)}"

@app.route('/', methods=['GET', 'POST'])
def chatbot():
    """Render chatbot interface."""
    if 'chat_history' not in session:
        session['chat_history'] = []  # Initialize chat history if not already present

    return render_template('index.html', chat_history=session['chat_history'])

@app.route('/ask', methods=['POST'])
def ask():
    """Handle user question via AJAX."""
    user_question = request.json['question']
    print(f"User question: {user_question}")  # Logging for debugging
    answer = get_gemini_answer(user_question)

    if 'chat_history' not in session:
        session['chat_history'] = [] 

    # Add the user question and bot answer to chat history
    session['chat_history'].append({'user': user_question, 'bot': answer})
    session.modified = True  # Mark session as modified for Flask

    print(f"Updated chat history in ask: {session['chat_history']}")  # Debugging information
    return jsonify({'answer': answer})


@app.route('/download/<format_type>', methods=['GET'])
def download_chat(format_type):
    """Download chat history in specified format."""
    if 'chat_history' not in session or not session['chat_history']:
        print("No chat history available in session")  
        return "No chat history available.", 400
    
    print(f"Chat history before download: {session['chat_history']}") 
    if format_type == 'pdf':
        return save_as_pdf(session['chat_history'])
    elif format_type == 'word':
        return save_as_word(session['chat_history'])
    elif format_type == 'txt':
        return save_as_txt(session['chat_history'])
    elif format_type == 'md':
        return save_as_md(session['chat_history'])
    return "Invalid format", 400

def save_as_pdf(chat_history):
    """Save chat history as PDF."""
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=12)

    for entry in chat_history:
        pdf.multi_cell(0, 10, f"You: {entry['user']}")
        pdf.multi_cell(0, 10, f"Chatbot: {entry['bot']}")
        pdf.cell(0, 10, "", ln=True)  

    filename = "chat_history.pdf"
    pdf.output(filename)
    return send_file(filename, as_attachment=True)

def save_as_word(chat_history):
    """Save chat history as Word document."""
    doc = Document()
    for entry in chat_history:
        doc.add_paragraph(f"You: {entry['user']}")
        doc.add_paragraph(f"Chatbot: {entry['bot']}")
        doc.add_paragraph() 

    filename = "chat_history.docx"
    doc.save(filename)
    return send_file(filename, as_attachment=True)

def save_as_txt(chat_history):
    """Save chat history as text file."""
    filename = "chat_history.txt"
    with open(filename, 'w') as file:
        for entry in chat_history:
            file.write(f"You: {entry['user']}\n")
            file.write(f"Chatbot: {entry['bot']}\n\n")  

    return send_file(filename, as_attachment=True)

def save_as_md(chat_history):
    """Save chat history as Markdown file."""
    filename = "chat_history.md"
    with open(filename, 'w') as file:
        file.write("# Chat History\n\n")
        for entry in chat_history:
            file.write(f"**You**: {entry['user']}\n\n")
            file.write(f"**Chatbot**: {entry['bot']}\n\n")
            file.write("---\n\n")  

    return send_file(filename, as_attachment=True)

@app.route('/chat-history', methods=['GET'])
def chat_history():
    """Return chat history as JSON."""
    if 'chat_history' in session:
        return jsonify({'chat_history': session['chat_history']})
    return jsonify({'chat_history': []})

if __name__ == '__main__':
    app.run(debug=True)
